from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUT = "/sessions/keen-sweet-fermi/mnt/outputs/Buildbox_SweetLand_Guide.docx"

doc = Document()

# ── Page setup (A4)
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = section.right_margin = Cm(2.5)
section.top_margin  = section.bottom_margin = Cm(2.5)

# ── Color palette
DARK_BG   = RGBColor(0x16, 0x21, 0x3E)
PINK      = RGBColor(0xF2, 0x9E, 0xBC)
TEAL      = RGBColor(0x7D, 0xD8, 0xC6)
GOLD      = RGBColor(0xF8, 0xD0, 0x40)
ORANGE    = RGBColor(0xFF, 0xD6, 0xA5)
PURPLE    = RGBColor(0xC3, 0xB1, 0xE1)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
DARK_TEXT = RGBColor(0x1A, 0x1A, 0x2E)
GRAY      = RGBColor(0x88, 0x88, 0x99)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_para(text, size=11, bold=False, color=DARK_TEXT, align=None, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if align: p.alignment = align
    run = p.add_run(text)
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.color.rgb = color
    return p

def add_heading(text, level=1):
    colors = {1: PINK, 2: TEAL, 3: GOLD}
    sizes  = {1: 20, 2: 15, 3: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level==1 else 10)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.size  = Pt(sizes[level])
    run.font.bold  = True
    run.font.color.rgb = colors.get(level, DARK_TEXT)
    return p

def add_step_table(steps):
    """Create a numbered step table with icon + description."""
    table = doc.add_table(rows=len(steps), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (title, desc) in enumerate(steps):
        row = table.rows[i]
        # Num cell
        c0 = row.cells[0]
        c0.width = Inches(0.5)
        set_cell_bg(c0, 'F29EBC')
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(str(i+1))
        r0.font.bold  = True
        r0.font.size  = Pt(14)
        r0.font.color.rgb = WHITE
        # Content cell
        c1 = row.cells[1]
        set_cell_bg(c1, 'F8F4FF')
        p1 = c1.paragraphs[0]
        r1a = p1.add_run(title + "  ")
        r1a.font.bold  = True
        r1a.font.size  = Pt(11)
        r1a.font.color.rgb = DARK_TEXT
        r1b = p1.add_run(desc)
        r1b.font.size  = Pt(10)
        r1b.font.color.rgb = GRAY
    doc.add_paragraph()

def add_info_box(label, text, bg='E8F8F4', label_color='7DD8C6'):
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    row = table.rows[0]
    c0 = row.cells[0]
    c0.width = Inches(1.2)
    set_cell_bg(c0, label_color.lstrip('#'))
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(label)
    r0.font.bold  = True
    r0.font.size  = Pt(10)
    r0.font.color.rgb = WHITE
    c1 = row.cells[1]
    set_cell_bg(c1, bg)
    p1 = c1.paragraphs[0]
    r1 = p1.add_run(text)
    r1.font.size  = Pt(10)
    r1.font.color.rgb = DARK_TEXT
    doc.add_paragraph()

def add_asset_table(rows_data, headers):
    table = doc.add_table(rows=1+len(rows_data), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        set_cell_bg(hrow.cells[i], '162130')
        p = hrow.cells[i].paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = WHITE
    # Data rows
    for ri, row_data in enumerate(rows_data):
        drow = table.rows[ri+1]
        bg = 'F8F4FF' if ri%2==0 else 'FFFFFF'
        for ci, cell_text in enumerate(row_data):
            set_cell_bg(drow.cells[ci], bg)
            p = drow.cells[ci].paragraphs[0]
            r = p.add_run(str(cell_text))
            r.font.size  = Pt(9)
            r.font.color.rgb = DARK_TEXT
    doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# COVER
# ════════════════════════════════════════════════════════════════════════════
add_para("", space_after=20)
add_para("🍭  SWEET LAND ADVENTURE", 26, bold=True, color=PINK,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=20, space_after=4)
add_para("Buildbox Classic 개발 가이드", 16, bold=True, color=TEAL,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para("에셋 임포트 · 씬 구성 · 물리 설정 · 빌드 완전 가이드", 11, color=GRAY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)
add_para("─" * 72, 9, color=PURPLE, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("", space_after=20)

# ════════════════════════════════════════════════════════════════════════════
# 0. 준비 사항
# ════════════════════════════════════════════════════════════════════════════
add_heading("0.  시작 전 체크리스트", 1)
add_info_box("📁 프로젝트 파일",
    "SweetLandAdventure.bbdoc + sweet_land_assets/ 폴더가 같은 위치에 있는지 확인하세요.",
    bg='FFF8E8', label_color='F8D040')
add_info_box("💾 에셋 폴더",
    "outputs\\sweet_land_assets\\ 안에 PNG 파일 24개가 있어야 합니다.",
    bg='E8F8F4', label_color='7DD8C6')
add_info_box("🎮 소프트웨어",
    "Buildbox Classic 2.x 가 설치되어 있어야 합니다 (바탕화면 아이콘 확인).",
    bg='F4E8F8', label_color='C3B1E1')

# ════════════════════════════════════════════════════════════════════════════
# 1. 프로젝트 열기
# ════════════════════════════════════════════════════════════════════════════
add_heading("1.  프로젝트 열기", 1)
add_para("Buildbox Classic을 실행하고 미리 생성된 .bbdoc 파일을 불러옵니다.", space_after=8)
add_step_table([
    ("Buildbox 실행",       "바탕화면의 'Buildbox Classic' 아이콘을 더블클릭합니다."),
    ("프로젝트 열기",        "상단 메뉴 File → Open Project 클릭합니다."),
    (".bbdoc 선택",         "outputs 폴더에서 SweetLandAdventure.bbdoc 선택 → 열기 클릭"),
    ("에셋 경로 확인",       "이미지가 보이지 않으면 Settings → Asset Path를 sweet_land_assets로 재설정"),
    ("저장",                "Ctrl+S 로 즉시 저장합니다. 작업 전 반드시 백업하세요."),
])
add_info_box("⚠️  주의",
    ".bbdoc 파일과 sweet_land_assets 폴더는 반드시 같은 디렉토리에 있어야 에셋이 정상 로드됩니다.",
    bg='FFF0F0', label_color='E05060')

# ════════════════════════════════════════════════════════════════════════════
# 2. 에셋 임포트
# ════════════════════════════════════════════════════════════════════════════
add_heading("2.  에셋 임포트 및 등록", 1)
add_para("Buildbox Classic에서 PNG를 임포트하는 방법은 두 가지입니다.", space_after=8)

add_heading("2-1. 드래그 앤 드롭 방식 (권장)", 2)
add_step_table([
    ("Asset Browser 열기",  "우측 패널에서 'Assets' 탭 클릭"),
    ("PNG 폴더 열기",        "파일 탐색기로 outputs\\sweet_land_assets\\ 폴더를 열어 두기"),
    ("드래그",              "PNG 파일을 Asset Browser로 끌어다 놓기"),
    ("카테고리 지정",        "각 에셋을 드래그 후 Sprites / UI / Background 중 선택"),
    ("확인",                "썸네일이 보이면 임포트 완료"),
])

add_heading("2-2. 에셋별 임포트 카테고리 표", 2)
add_asset_table([
    ["cookie_idle.png",          "Sprite",     "주인공 대기 프레임"],
    ["cookie_run1~4.png",        "Sprite",     "달리기 애니메이션 4장"],
    ["cookie_jump.png",          "Sprite",     "점프 프레임"],
    ["enemy_donut.png",          "Sprite",     "도넛 롤러 적"],
    ["enemy_lollipop_fairy.png", "Sprite",     "롤리팝 요정 적"],
    ["enemy_caramel_blob.png",   "Sprite",     "캐러멜 블롭 적"],
    ["tile_cookie_platform.png", "Tile",       "기본 발판 타일"],
    ["tile_marshmallow.png",     "Tile",       "바운스 발판 타일"],
    ["tile_cotton_candy.png",    "Tile",       "솜사탕 발판 타일"],
    ["tile_item_block.png",      "Tile",       "아이템 블록"],
    ["item_candy_coin.png",      "Collectible","캔디 코인 (점수+100)"],
    ["item_sugar_star.png",      "Collectible","슈거 스타 (점수+1000)"],
    ["item_cake_slice.png",      "Collectible","케이크 조각 (HP+1)"],
    ["item_rainbow_candy.png",   "Power-up",   "레인보우 캔디 (무적 10초)"],
    ["ui_heart_full.png",        "UI",         "체력 하트 (풀)"],
    ["ui_heart_empty.png",       "UI",         "체력 하트 (빈)"],
    ["prop_cupcake_house.png",   "Background", "컵케이크 건물 오브젝트"],
    ["prop_lollipop_tree.png",   "Background", "롤리팝 나무 오브젝트"],
    ["prop_goal_flag.png",       "Goal",       "스테이지 목표 깃발"],
    ["bg_cloud.png",             "Background", "구름 배경 요소"],
    ["bg_cotton_tree.png",       "Background", "솜사탕 나무 배경"],
], ["파일명", "카테고리", "역할"])

# ════════════════════════════════════════════════════════════════════════════
# 3. 캐릭터 설정
# ════════════════════════════════════════════════════════════════════════════
add_heading("3.  캐릭터 설정", 1)
add_para("상단 메뉴 Game → Characters에서 쿠키 캐릭터를 설정합니다.", space_after=8)

add_heading("3-1. 기본 속성 값", 2)
add_asset_table([
    ["이름",         "Cookie"],
    ["크기",         "32 × 32 pixels"],
    ["이동 속도",     "250"],
    ["점프 힘",       "500  (가변 점프: Hold to Jump 활성화)"],
    ["최대 체력",     "5"],
    ["목숨 수",       "3"],
    ["중력 배율",     "1.0"],
    ["충돌 박스",     "Box  /  오프셋 0,0  /  크기 28×30"],
], ["속성", "값"])

add_heading("3-2. 애니메이션 등록", 2)
add_step_table([
    ("캐릭터 선택",    "Characters 목록에서 Cookie 선택"),
    ("Idle 등록",     "Animation 탭 → Add → 'Idle' 이름 입력 → cookie_idle.png 지정, FPS: 4, Loop: ON"),
    ("Run 등록",      "Add → 'Run' → cookie_run1~4.png 순서대로 추가, FPS: 8, Loop: ON"),
    ("Jump 등록",     "Add → 'Jump' → cookie_jump.png, FPS: 6, Loop: OFF"),
    ("상태 연결",      "State 탭 → Idle/Run/Jump 각각 위 애니메이션으로 연결"),
])

# ════════════════════════════════════════════════════════════════════════════
# 4. 월드 & 씬 구성
# ════════════════════════════════════════════════════════════════════════════
add_heading("4.  월드 & 씬 구성", 1)
add_para("이미 .bbdoc에 World 1 기본 구조가 담겨 있습니다. 아래 값으로 조정하세요.", space_after=8)

add_heading("4-1. 씬 설정 (Scene Properties)", 2)
add_asset_table([
    ["씬 크기",      "1920 × 216 pixels"],
    ["스크롤 방향",   "Horizontal (좌→우)"],
    ["배경 색 (상)",  "#7DD8C6  (민트 틸)"],
    ["배경 색 (하)",  "#FAD5B5  (크림 피치)"],
    ["제한 시간",     "300초 (보스 씬은 0 = 무제한)"],
    ["다음 씬",      "nextScene 필드에 다음 씬 ID 입력"],
], ["항목", "값"])

add_heading("4-2. 발판 배치 기준", 2)
add_asset_table([
    ["쿠키 발판",    "Y=32  (바닥), 너비 타일링, Solid=ON"],
    ["마시멜로 발판", "Y=112~144, Bounce Force=600"],
    ["솜사탕 발판",  "Y=112~144, Disappear=ON, 3초"],
    ["이동 발판",    "Move Type=Patrol, 좌우 이동 범위 설정"],
    ["아이템 블록",  "발판 위 Y=144~160, 아래에서 충돌 시 아이템 팝"],
], ["타입", "배치 가이드"])

add_heading("4-3. 배경 패럴랙스 설정", 2)
add_step_table([
    ("레이어 -2 (구름)",    "bg_cloud.png → Layer: -2, Scroll Speed X: 0.3, Tile: ON"),
    ("레이어 -1 (나무/건물)", "prop_*.png → Layer: -1, Scroll Speed X: 0.6"),
    ("레이어 0 (발판/적)",   "tile_*.png, enemy_*.png → Layer: 0"),
    ("레이어 1 (아이템/HUD)", "item_*.png, ui_*.png → Layer: 1"),
])

# ════════════════════════════════════════════════════════════════════════════
# 5. 물리 & 게임플레이 설정
# ════════════════════════════════════════════════════════════════════════════
add_heading("5.  물리 & 게임플레이 설정", 1)
add_heading("5-1. Physics 값", 2)
add_asset_table([
    ["중력 (Gravity)",    "-20"],
    ["공기 저항",          "0.95"],
    ["마찰력",             "0.85  (솜사탕 발판: 0.3으로 낮게)"],
    ["점프 방식",          "Hold to Jump = ON  (가변 점프 높이 구현)"],
    ["낙하 가속",          "Fall Multiplier: 1.5  (낙하감 개선)"],
], ["항목", "값"])

add_heading("5-2. 적 AI 설정", 2)
add_asset_table([
    ["도넛 롤러",          "Patrol AI, Speed: 80, Kill On Jump: ON, Damage: 1"],
    ["롤리팝 요정",        "Flying AI, Y 진동, Speed: 60, Kill On Dash: ON"],
    ["캐러멜 블롭",        "Patrol AI, HP: 3, Speed: 40, Jump-kill 3회"],
], ["적 이름", "AI 설정"])

add_heading("5-3. 수집 아이템 값", 2)
add_asset_table([
    ["캔디 코인",          "+100점, 100개 수집 시 목숨+1"],
    ["슈거 스타",          "+1000점, 히든 아이템 (tag=star)"],
    ["케이크 조각",        "HP +1 회복"],
    ["레인보우 캔디",      "무적 10초 + 이동속도 2배"],
], ["아이템", "효과"])

# ════════════════════════════════════════════════════════════════════════════
# 6. HUD 설정
# ════════════════════════════════════════════════════════════════════════════
add_heading("6.  HUD (인게임 UI) 설정", 1)
add_step_table([
    ("하트 (체력)",   "ui_heart_full/empty.png → HUD Editor → 좌상단 배치, 최대 5개, 간격 20px"),
    ("코인 카운터",   "item_candy_coin.png 아이콘 + 숫자 텍스트 → 우상단"),
    ("점수",          "Score 텍스트 위젯 → 중앙 상단, 폰트 크기 16"),
    ("스테이지",      "World/Scene 번호 텍스트 → 중앙 상단"),
    ("타이머",        "Timer 위젯 → 우상단 코인 아래, 빨간색"),
])

# ════════════════════════════════════════════════════════════════════════════
# 7. 스코어링 시스템
# ════════════════════════════════════════════════════════════════════════════
add_heading("7.  스코어링 시스템", 1)
add_asset_table([
    ["코인 수집",          "+100점"],
    ["적 밟기 처치",        "+200점"],
    ["연속 처치 콤보",      "×1.5 배율 (콤보 2회 이상 시)"],
    ["슈거 스타 획득",      "+1,000점"],
    ["스테이지 클리어",     "잔여 시간 × 10점 보너스"],
    ["무피해 클리어",       "+5,000점 보너스"],
    ["High Score 저장",    "Game Settings → Score → Enable High Score: ON"],
], ["이벤트", "점수"])

# ════════════════════════════════════════════════════════════════════════════
# 8. 빌드 & 테스트
# ════════════════════════════════════════════════════════════════════════════
add_heading("8.  빌드 & 테스트", 1)
add_heading("8-1. 에디터 내 플레이 테스트", 2)
add_step_table([
    ("Play 버튼",           "에디터 상단 ▶ Play 클릭 또는 Space 키"),
    ("캐릭터 조작 확인",     "←→ 이동, ↑ 점프, X 대시 입력 테스트"),
    ("충돌 확인",            "발판에 정상 착지, 적에게 피해, 코인 수집 확인"),
    ("씬 전환 확인",         "목표 깃발에 도달 시 다음 씬으로 이동 확인"),
    ("Stop",                "Esc 또는 Stop 버튼으로 플레이 종료"),
])

add_heading("8-2. 내보내기 (Export)", 2)
add_step_table([
    ("Export 메뉴",         "File → Export → 플랫폼 선택 (Windows / Android / iOS)"),
    ("설정 확인",           "게임 이름: SweetLandAdventure, 버전: 1.0, 아이콘 설정"),
    ("빌드",                "Export 버튼 클릭 → 빌드 완료까지 대기"),
    ("테스트",              "생성된 .exe 또는 .apk 실행하여 최종 확인"),
])

add_info_box("💡 팁",
    "Ctrl+Z 실행 취소 / Ctrl+S 수시로 저장 / 씬 수정 후 반드시 플레이 테스트로 확인하세요.",
    bg='E8F8F4', label_color='7DD8C6')

# ════════════════════════════════════════════════════════════════════════════
# 9. 트러블슈팅
# ════════════════════════════════════════════════════════════════════════════
add_heading("9.  자주 묻는 트러블슈팅", 1)
add_asset_table([
    ["에셋이 핑크/체크로 보임",   "sweet_land_assets 폴더가 .bbdoc과 같은 위치인지 확인"],
    ["캐릭터가 떨어짐",           "Physics → Gravity 값이 -20인지, 발판 Solid=ON인지 확인"],
    ["애니메이션이 안 바뀜",      "Character → State 탭에서 Idle/Run/Jump 상태 연결 재확인"],
    ["적이 움직이지 않음",        "Enemy Properties → Patrol Left/Right 범위 값이 있는지 확인"],
    ["코인 수집이 안 됨",         "Collectible Type이 coin인지, 충돌 레이어 확인"],
    ["다음 씬으로 안 넘어감",     "Goal 오브젝트 → triggerNextScene=true, nextScene ID 확인"],
    [".bbdoc 로드 실패",          "Buildbox Classic 버전이 2.x인지 확인. v3은 다른 포맷 사용"],
], ["증상", "해결 방법"])

# ════════════════════════════════════════════════════════════════════════════
# 10. 에셋 파일 목록
# ════════════════════════════════════════════════════════════════════════════
add_heading("10.  제공 에셋 전체 목록", 1)
add_para("sweet_land_assets/ 폴더에 아래 24개 PNG 파일이 포함되어 있습니다. 모든 파일은 원본 해상도의 4× 업스케일 (픽셀퍼펙트 NEAREST) PNG입니다.", space_after=8)
add_asset_table([
    ["캐릭터",  "cookie_idle.png, cookie_run1~4.png, cookie_jump.png  (6개)"],
    ["적",      "enemy_donut.png, enemy_lollipop_fairy.png, enemy_caramel_blob.png  (3개)"],
    ["타일",    "tile_cookie_platform.png, tile_marshmallow.png, tile_cotton_candy.png, tile_item_block.png  (4개)"],
    ["아이템",  "item_candy_coin.png, item_sugar_star.png, item_cake_slice.png, item_rainbow_candy.png  (4개)"],
    ["UI",      "ui_heart_full.png, ui_heart_empty.png  (2개)"],
    ["배경/소품", "prop_cupcake_house.png, prop_lollipop_tree.png, prop_goal_flag.png, bg_cloud.png, bg_cotton_tree.png  (5개)"],
], ["카테고리", "파일 목록"])

# Footer
add_para("")
add_para("─" * 72, 9, color=PURPLE)
add_para("Sweet Land Adventure  ·  Buildbox Classic 가이드  v1.0  ·  Claude 생성",
         9, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.save(OUT)
print(f"✅ 가이드 문서 생성 완료: {OUT}")
